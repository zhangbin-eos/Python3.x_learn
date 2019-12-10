
# -*- coding: utf-8 -*-

import sys

import docx
from docx.shared import Inches

import json


def main(argv=None):
    document = docx.Document()
    document.add_heading('标题', 0)  #插入标题
    document.add_page_break()

    document.add_heading('标题1', level=1)
    document.add_heading('标题2', level=2)
    document.add_heading('标题3', level=3)
    document.add_heading('标题4', level=4)
    document.add_heading('标题5', level=5)
    document.add_heading('标题6', level=6)
    print("document",type(document))
    paragraph = document.add_paragraph('一个段落, ')   #插入段落

    print("paragraph",type(paragraph))
    paragraph.add_run('正常')#==paragraph.add_run().text="'正常'"

    print("paragraph.add_run()",type(paragraph.add_run()))
    
    paragraph.add_run().text=str(help(paragraph.add_run()))

    paragraph.add_run('粗体').bold = True
    paragraph.add_run('斜体').italic = True
    paragraph.add_run('下划线').underline = True
    
    #document.add_picture(*.png', width=Inches(1.25)) #插入图片

    table = document.add_table(rows=3, cols=3,style="Table Grid") #插入表格
    print("table",type(table))
    hdr_cells = table.rows[0].cells
    print("hdr_cells",type(hdr_cells))
    hdr_cells[0].text = '0'
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    document.add_page_break()

    document.save('demo.docx')  #保存文档


if __name__ == "__main__":
    sys.exit(main())
